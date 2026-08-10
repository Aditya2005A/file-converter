"""
Automated verification script for the File Toolkit Flask app.

Runs entirely against Flask's test_client() — no external server or network
access is required. Exercises every API endpoint with realistic in-memory
payloads and asserts on status codes, content types, and byte signatures.
"""

import io
import sys

from PIL import Image
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

sys.path.insert(0, "api")
from index import app  # noqa: E402


def make_dummy_image_bytes(fmt="PNG", size=(64, 64), color=(255, 0, 0)):
    buf = io.BytesIO()
    img = Image.new("RGB", size, color=color)
    img.save(buf, format=fmt)
    buf.seek(0)
    return buf


def make_dummy_docx_bytes():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("This is a test paragraph for the file toolkit.")
    doc.add_paragraph("Second paragraph with some more sample text content.")
    doc.save(buf)
    buf.seek(0)
    return buf


def make_dummy_pdf_bytes():
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, 720, "Dummy PDF for rotation testing.")
    c.save()
    buf.seek(0)
    return buf


def run_tests():
    app.config["TESTING"] = True
    client = app.test_client()

    results = []

    # -----------------------------------------------------------------
    # Test 1: Health check
    # -----------------------------------------------------------------
    try:
        resp = client.get("/api/health")
        assert resp.status_code == 200, f"expected 200, got {resp.status_code}"
        assert resp.get_json() == {"status": "ok"}, f"unexpected body: {resp.get_json()}"
        results.append(("Test 1: GET /api/health", True, ""))
    except Exception as exc:
        results.append(("Test 1: GET /api/health", False, str(exc)))

    # -----------------------------------------------------------------
    # Test 2: Image conversion (PNG -> WEBP, quality=80, rotate=90)
    # -----------------------------------------------------------------
    try:
        img_buf = make_dummy_image_bytes(fmt="PNG", size=(100, 50))
        data = {
            "file": (img_buf, "test.png"),
            "format": "WEBP",
            "quality": "80",
            "rotate": "90",
        }
        resp = client.post(
            "/api/convert-image",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 200, f"expected 200, got {resp.status_code}: {resp.data[:300]}"
        assert resp.mimetype == "image/webp", f"unexpected mimetype: {resp.mimetype}"

        out_img = Image.open(io.BytesIO(resp.data))
        assert out_img.format == "WEBP", f"unexpected output format: {out_img.format}"
        # Original was 100x50; a 90-degree rotation with expand=True should swap dims.
        assert out_img.size == (50, 100), f"unexpected size after rotation: {out_img.size}"

        results.append(("Test 2: POST /api/convert-image (WEBP, q=80, rotate=90)", True, ""))
    except Exception as exc:
        results.append(("Test 2: POST /api/convert-image (WEBP, q=80, rotate=90)", False, str(exc)))

    # -----------------------------------------------------------------
    # Test 3: File-to-PDF (TXT + DOCX combined)
    # -----------------------------------------------------------------
    try:
        txt_buf = io.BytesIO(b"Hello from a plain text file.\nSecond line of text.\n")
        docx_buf = make_dummy_docx_bytes()

        data = {
            "file": [
                (txt_buf, "sample.txt"),
                (docx_buf, "sample.docx"),
            ]
        }
        resp = client.post(
            "/api/to-pdf",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 200, f"expected 200, got {resp.status_code}: {resp.data[:300]}"
        assert resp.mimetype == "application/pdf", f"unexpected mimetype: {resp.mimetype}"
        assert resp.data[:5] == b"%PDF-", "response does not start with %PDF- header"

        results.append(("Test 3: POST /api/to-pdf (TXT + DOCX)", True, ""))
    except Exception as exc:
        results.append(("Test 3: POST /api/to-pdf (TXT + DOCX)", False, str(exc)))

    # -----------------------------------------------------------------
    # Test 3b (bonus): File-to-PDF with an image input too
    # -----------------------------------------------------------------
    try:
        img_buf = make_dummy_image_bytes(fmt="JPEG", size=(80, 80))
        data = {"file": (img_buf, "sample.jpg")}
        resp = client.post(
            "/api/to-pdf",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 200, f"expected 200, got {resp.status_code}: {resp.data[:300]}"
        assert resp.data[:5] == b"%PDF-", "response does not start with %PDF- header"

        results.append(("Test 3b: POST /api/to-pdf (image input)", True, ""))
    except Exception as exc:
        results.append(("Test 3b: POST /api/to-pdf (image input)", False, str(exc)))

    # -----------------------------------------------------------------
    # Test 4: PDF rotation
    # -----------------------------------------------------------------
    try:
        pdf_buf = make_dummy_pdf_bytes()
        data = {"file": (pdf_buf, "sample.pdf"), "angle": "90"}
        resp = client.post(
            "/api/rotate-pdf",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 200, f"expected 200, got {resp.status_code}: {resp.data[:300]}"
        assert resp.mimetype == "application/pdf", f"unexpected mimetype: {resp.mimetype}"
        assert resp.data[:5] == b"%PDF-", "response does not start with %PDF- header"

        results.append(("Test 4: POST /api/rotate-pdf", True, ""))
    except Exception as exc:
        results.append(("Test 4: POST /api/rotate-pdf", False, str(exc)))

    # -----------------------------------------------------------------
    # Report
    # -----------------------------------------------------------------
    print("\n" + "=" * 60)
    print("TEST RESULTS")
    print("=" * 60)
    all_passed = True
    for name, passed, detail in results:
        status = "PASS" if passed else "FAIL"
        print(f"[{status}] {name}")
        if not passed:
            print(f"       -> {detail}")
            all_passed = False
    print("=" * 60)

    if all_passed:
        print(f"ALL {len(results)} TESTS PASSED ✅")
    else:
        failed_count = sum(1 for _, p, _ in results if not p)
        print(f"{failed_count} of {len(results)} TESTS FAILED ❌")
    print("=" * 60 + "\n")

    return all_passed


if __name__ == "__main__":
    success = run_tests()
    sys.exit(0 if success else 1)
